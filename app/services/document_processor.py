"""
Legal Document Review System - Document Processor Service
Handles PDF, DOCX, and TXT file parsing and text extraction.
"""
import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pdfplumber


class DocumentProcessor:
    """Service for processing and extracting text from documents."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    def __init__(self, upload_dir: str = "storage/uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    async def save_file(self, filename: str, content: bytes) -> Path:
        """Save uploaded file to disk."""
        file_path = self.upload_dir / filename
        with open(file_path, "wb") as f:
            f.write(content)
        return file_path
    
    def get_file_extension(self, filename: str) -> str:
        """Get file extension from filename."""
        return Path(filename).suffix.lower()
    
    def validate_file(self, filename: str) -> bool:
        """Validate if file type is supported."""
        ext = self.get_file_extension(filename)
        return ext in self.SUPPORTED_EXTENSIONS
    
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from document based on file type."""
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return await self._extract_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return await self._extract_docx(file_path)
        elif ext == '.txt':
            return await self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    async def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using pdfplumber for better accuracy."""
        text_parts = []
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception:
            # Fallback to PyPDF2
            reader = PdfReader(str(file_path))
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    async def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(str(file_path))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    async def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
        """Split text into overlapping chunks for embedding."""
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence-ending punctuation
                for i in range(end, max(start + chunk_size // 2, start), -1):
                    if text[i] in '.!?\n':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap if end < text_length else text_length
        
        return chunks
    
    async def cleanup_file(self, file_path: Path) -> None:
        """Remove temporary file after processing."""
        if file_path.exists():
            os.remove(file_path)
